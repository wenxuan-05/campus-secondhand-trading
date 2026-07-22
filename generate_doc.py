#!/usr/bin/env python3
"""生成校园二手交易平台——功能概述与实现步骤 Word 文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import datetime

doc = Document()

# ── 页面设置 ──
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── 样式定义 ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.25

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_font = heading_style.font
    heading_font.name = '微软雅黑'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    heading_font.color.rgb = RGBColor(0x1A, 0x56, 0xDB) if level <= 2 else RGBColor(0x33, 0x33, 0x33)
    if level == 1:
        heading_font.size = Pt(20)
    elif level == 2:
        heading_font.size = Pt(14)
    else:
        heading_font.size = Pt(12)


def add_para(text, bold=False, size=None, color=None, align=None, indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*color)
    if align: p.alignment = align
    if indent: p.paragraph_format.left_indent = Cm(indent)
    return p


def add_table(headers, rows, col_widths=None):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Light Grid Accent 1')
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8.5)
    doc.add_paragraph()  # 空行


# ══════════════════════════════════════════
# 封面
# ══════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

add_para('校园二手交易平台', bold=True, size=32, color=(0x1A, 0x56, 0xDB), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('「淘鸭」', bold=True, size=22, color=(0xFF, 0xD0, 0x00), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('', size=12)
add_para('功能概述与实现步骤', bold=True, size=18, color=(0x55, 0x55, 0x55), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('', size=12)
add_para(f'文档版本：V1.0', size=10, color=(0x88, 0x88, 0x88), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(f'生成日期：{datetime.date.today().strftime("%Y年%m月%d日")}', size=10, color=(0x88, 0x88, 0x88), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('技术栈：Spring Boot 3 + Vue 3 + MyBatis-Plus + SQL Server + Redis + MinIO', size=9, color=(0xAA, 0xAA, 0xAA), align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ══════════════════════════════════════════
# 目录页
# ══════════════════════════════════════════
doc.add_heading('目  录', level=1)
toc_items = [
    '一、项目概述 ............................................. 3',
    '二、技术架构 ............................................. 3',
    '三、系统功能模块总览 ....................................... 4',
    '四、核心功能详解 .......................................... 5',
    '    4.1 用户模块 .......................................... 5',
    '    4.2 商品模块 .......................................... 5',
    '    4.3 订单交易模块 ....................................... 6',
    '    4.4 即时通讯模块 ....................................... 6',
    '    4.5 收藏模块 ........................................... 6',
    '    4.6 评价模块 ........................................... 7',
    '    4.7 退款/售后模块 ...................................... 7',
    '    4.8 求购广场模块 ....................................... 7',
    '    4.9 校园大使模块 ....................................... 7',
    '    4.10 举报管理模块 ...................................... 7',
    '    4.11 个性化推荐模块 .................................... 8',
    '    4.12 管理后台模块 ...................................... 8',
    '五、数据库设计 ............................................ 8',
    '六、安全性设计 ............................................ 9',
    '七、部署架构 .............................................. 9',
    '八、实现步骤 .............................................. 10',
]
for item in toc_items:
    add_para(item, size=10)

doc.add_page_break()

# ══════════════════════════════════════════
# 一、项目概述
# ══════════════════════════════════════════
doc.add_heading('一、项目概述', level=1)

add_para('"淘鸭"是一个专为西南财经大学（SWUFE）在校学生打造的校园二手商品交易平台，旨在为同学们提供安全、便捷、高效的闲置物品买卖服务。平台以"鸭鸭黄"（#FFD000）为主题色，整体设计风格亲切活泼，贴近校园生活。')

add_para('核心业务场景：卖家发布闲置商品 → 平台审核（校园大使） → 买家浏览/下单/支付 → 卖家生成取货码 → 双方面对面验货交付 → 确认收货并互评。整个交易闭环在线下交付环节通过6位数字取货码保证安全性。')

add_para('平台特色：', bold=True)
features = [
    '校园大使审核机制：招募本校学生担任"校园大使"，负责商品内容审核，保证平台内容质量；',
    '信用评分体系：用户交易行为（确认收货速度、评价质量、退款率等）影响信用分，高信用用户享受更多权益；',
    '双角色后台：管理员（ADMIN）和校园大使（CAMPUS_AMBASSADOR）拥有不同权限的后台面板；',
    '实时即时通讯：WebSocket 实现买卖双方在线沟通，内置 DFA 敏感词过滤；',
    '个性化推荐：基于用户浏览/收藏/购买行为的协同过滤推荐算法；',
    '线下取货码验证：6位数字取货码保障面对面交易安全，防止冒领。',
]
for f in features:
    p = doc.add_paragraph(f, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)

# ══════════════════════════════════════════
# 二、技术架构
# ══════════════════════════════════════════
doc.add_heading('二、技术架构', level=1)

add_para('2.1 技术选型', bold=True, size=12)
add_table(
    ['层级', '技术', '版本', '说明'],
    [
        ['后端框架', 'Spring Boot', '3.2.0', '提供 RESTful API 服务'],
        ['ORM 框架', 'MyBatis-Plus', '3.5.5', '简化 CRUD，Lambda 查询'],
        ['数据库', 'SQL Server', '2022', '通过 mssql-jdbc 连接'],
        ['缓存', 'Redis (Lettuce)', '7.x', 'Token 缓存、验证码存储'],
        ['认证', 'JWT (jjwt)', '0.12.3', 'HMAC-SHA384 签名，7天有效期'],
        ['文件存储', 'MinIO', '最新版', '兼容 S3 协议的对象存储'],
        ['邮件服务', 'Spring Mail', '3.2.0', 'QQ邮箱 SMTP 发送验证码'],
        ['即时通讯', 'WebSocket', '—', 'Spring WebSocket 原生支持'],
        ['前端框架', 'Vue 3', '3.4.x', 'Composition API + <script setup>'],
        ['构建工具', 'Vite', '5.1.x', '快速 HMR 开发体验'],
        ['UI 组件库', 'Element Plus', '2.5.x', '企业级中后台组件'],
        ['状态管理', 'Pinia', '2.1.7', 'Vue 3 官方推荐'],
        ['路由', 'Vue Router', '4.3.x', 'History 模式，路由守卫'],
        ['HTTP 客户端', 'Axios', '1.6.x', '请求/响应拦截器'],
    ],
    col_widths=[2.5, 2.5, 1.5, 5.5]
)

add_para('2.2 项目结构', bold=True, size=12)
add_table(
    ['目录', '说明'],
    [
        ['backend/src/main/java/.../controller/', '控制器层，12个 Controller 处理 80+ API'],
        ['backend/src/main/java/.../service/', '服务层接口 + impl 实现'],
        ['backend/src/main/java/.../mapper/', 'MyBatis-Plus Mapper 接口（含 XML）'],
        ['backend/src/main/java/.../entity/', '实体类：User, Goods, Order, ChatMessage 等 11 张表'],
        ['backend/src/main/java/.../dto/', '数据传输对象：请求/响应 DTO + VO'],
        ['backend/src/main/java/.../config/', '配置类：CORS, MinIO, WebSocket, MyBatis-Plus, 定时任务'],
        ['backend/src/main/java/.../security/', '安全：JWT Token Provider, UserContext, @RequireRole'],
        ['backend/src/main/java/.../interceptor/', '拦截器：WebSocket 认证, 角色权限'],
        ['backend/src/main/java/.../websocket/', 'WebSocket Handler：ChatWebSocketHandler'],
        ['backend/src/main/java/.../common/', '公共类：Result, BusinessException, GlobalExceptionHandler'],
        ['admin/src/views/', 'Vue 视图页面（20+ 页面组件）'],
        ['admin/src/api/', 'Axios API 封装模块（8 个模块）'],
        ['admin/src/stores/', 'Pinia 状态管理（user, chat）'],
        ['admin/src/router/', 'Vue Router 路由配置（含权限守卫）'],
        ['admin/src/utils/', '工具函数（request 拦截器, format）'],
    ]
)

# ══════════════════════════════════════════
# 三、系统功能模块总览
# ══════════════════════════════════════════
doc.add_heading('三、系统功能模块总览', level=1)

add_table(
    ['序号', '功能模块', '核心功能', '面向角色'],
    [
        ['1', '用户模块', '注册/登录/找回密码/个人资料/信用分', '全部用户'],
        ['2', '商品模块', '发布/编辑/下架/搜索/浏览/刷新', '卖家/买家'],
        ['3', '订单交易模块', '下单/支付/取货码/验货/确认收货/取消', '买家/卖家'],
        ['4', '即时通讯模块', 'WebSocket 实时聊天/会话管理/图片发送', '买家/卖家'],
        ['5', '收藏模块', '添加/取消收藏/收藏列表', '买家'],
        ['6', '评价模块', '交易评价/评分/评价列表', '买家/卖家'],
        ['7', '退款/售后模块', '申请退款/商家处理/平台仲裁', '买家/卖家/管理员'],
        ['8', '求购广场模块', '发布求购/搜索求购/求购管理', '买家'],
        ['9', '校园大使模块', '申请大使/商品审核/推广统计', '校园大使/管理员'],
        ['10', '举报管理模块', '举报商品-用户/处理举报', '全部用户/管理员'],
        ['11', '个性化推荐模块', '智能推荐/关联推荐/行为记录', '买家'],
        ['12', '管理后台模块', '用户/商品/订单/退款管理/数据看板', '管理员/校园大使'],
        ['13', '文件服务模块', '图片上传/MinIO存储/图片代理', '全部用户'],
    ]
)

# ══════════════════════════════════════════
# 四、核心功能详解
# ══════════════════════════════════════════
doc.add_heading('四、核心功能详解', level=1)

# 4.1 用户模块
doc.add_heading('4.1 用户模块', level=2)
add_para('用户模块是整个系统的入口，提供完整的账号生命周期管理。')
add_table(
    ['接口', '方法', '路径', '说明'],
    [
        ['注册', 'POST', '/api/user/register', '学号+手机号+密码注册，密码 MD5 加密'],
        ['登录', 'POST', '/api/user/login', '返回 JWT Token，支持学号/手机号登录'],
        ['发送验证码', 'POST', '/api/user/send-code', 'QQ邮箱 SMTP 发送6位数字验证码'],
        ['重置密码', 'POST', '/api/user/reset-password', '通过邮箱验证码重置密码'],
        ['查看个人资料', 'GET', '/api/user/profile', '获取当前登录用户完整信息'],
        ['修改个人资料', 'PUT', '/api/user/profile', '修改昵称、头像、宿舍楼栋等'],
        ['用户公开信息', 'GET', '/api/user/{id}/public-profile', '查看用户公开名片（信用等级/交易次数/好评率）'],
        ['用户商品橱窗', 'GET', '/api/user/{id}/products', '查看某用户所有在售商品'],
    ]
)
add_para('用户角色体系：系统采用三级角色 RBAC 模型——普通用户（USER）、校园大使（CAMPUS_AMBASSADOR）、管理员（ADMIN）。管理员可手动提升用户为校园大使并分配工号，也可降级。校园大使可通过公开申请页面提交资料，经管理员审核通过后自动变更角色。')

# 4.2 商品模块
doc.add_heading('4.2 商品模块', level=2)
add_para('商品模块是平台的核心业务模块，支持完整的商品生命周期管理。商品状态流转：草稿(0) → 审核中(1) → 在售(2) → 已售出(3) → 已下架(4) → 审核失败(5) → 已删除(6)。')
add_table(
    ['操作', '方法', '路径', '说明'],
    [
        ['发布商品', 'POST', '/api/goods', '提交商品信息→状态变为"审核中(1)"'],
        ['编辑商品', 'PUT', '/api/goods/{id}', '修改商品信息'],
        ['商品下架', 'PUT', '/api/goods/{id}/off', '卖家主动下架，状态→4'],
        ['重新上架', 'PUT', '/api/goods/{id}/on', '下架商品重新上架，状态→2'],
        ['刷新商品', 'PUT', '/api/goods/{id}/refresh', '刷新排序权重（提高曝光）'],
        ['删除商品', 'DELETE', '/api/goods/{id}', '逻辑删除，状态→6'],
        ['商品详情', 'GET', '/api/goods/{id}', '返回商品完整信息'],
        ['商品搜索', 'GET', '/api/goods/search', '关键词/分类/价格/排序多条件搜索'],
        ['我的商品', 'GET', '/api/goods/my', '分页查看自己发布的商品'],
    ]
)
add_para('商品字段包括：标题、描述、价格、原价、分类、图片列表（多图）、视频链接、ISBN（书籍类）、成色（全新/几乎全新/良好/一般）、交易地点标签、是否可议价、是否支持邮寄、浏览/收藏/咨询次数、过期时间（发布后30天自动过期）、审核状态与审核意见等。')

# 4.3 订单交易模块
doc.add_heading('4.3 订单交易模块', level=2)
add_para('订单交易是平台最核心的业务流程，采用"线上支付 + 线下取货码验证"的模式。')
add_para('完整交易流程：', bold=True)
steps = [
    '买家点击"立即购买" → 创建待支付订单（status=1）；',
    '买家确认支付 → 订单状态变更为"已支付"(status=2)；',
    '卖家在订单详情中点击"生成取货码" → 系统生成6位随机数字码（status=3）；',
    '买卖双方约定线下见面，买家出示取货码，卖家输入验证 → 验证通过（status=4）；',
    '双方面对面查验商品无误后交付；',
    '买家点击"确认收货" → 订单完成（status=5）；',
    '买卖双方可互相评价。',
]
for s in steps:
    p = doc.add_paragraph(s, style='List Number')
    for run in p.runs:
        run.font.size = Pt(10)

add_para('买家可在支付前随时取消订单；取货码生成后买家也可取消，系统将自动进入退款流程。订单支持按"我是买家"和"我是卖家"两个维度分别查看。')

# 4.4 即时通讯模块
doc.add_heading('4.4 即时通讯模块', level=2)
add_para('即时通讯基于 Spring WebSocket 实现，买卖双方可以通过商品详情页直接发起对话。')
add_table(
    ['功能', '实现方式'],
    [
        ['WebSocket 连接', 'ChatWebSocketHandler，路径 /ws/chat，携带 JWT Token 认证'],
        ['会话管理', 'ChatSession 表记录会话双方，自动创建/查找已有会话'],
        ['消息发送', 'JSON 格式：{sessionId, receiverId, content, type}'],
        ['消息持久化', 'ChatMessage 表存储所有消息，支持分页拉取历史'],
        ['已读状态', '标记消息已读，会话列表展示未读数'],
        ['图片消息', '先上传到 MinIO (chat/ 目录)，发送图片 URL'],
        ['敏感词过滤', 'DFA 字典树算法实时过滤消息中的敏感词，替换为 ***'],
    ]
)

# 4.5 收藏模块
doc.add_heading('4.5 收藏模块', level=2)
add_para('买家可收藏感兴趣的商品，系统会检查是否已收藏（防止重复收藏），收藏列表支持分页查询。收藏数据同时作为推荐算法的用户行为输入之一。')

# 4.6 评价模块
doc.add_heading('4.6 评价模块', level=2)
add_para('交易完成后买卖双方可相互评价，评价包含评分（1-5星）和文字评论。评价与具体订单关联，系统会聚合计分到用户公开资料中的"好评率"和"交易次数"。评价创建后不可修改。')

# 4.7 退款/售后模块
doc.add_heading('4.7 退款/售后模块', level=2)
add_para('买家可在交易完成后申请退款，退款流程如下：')
refund_steps = [
    '买家提交退款申请 → 填写原因、上传凭证图片（支持多图，存为 JSON 数组）；',
    '退款状态流转：待卖家处理(1) → 卖家同意(2)/卖家拒绝(3)；',
    '卖家拒绝后：买家可申诉 → 管理员仲裁 → 平台介入(4)；',
    '管理员审核 → 仲裁通过(5)/仲裁驳回(6)。',
]
for s in refund_steps:
    p = doc.add_paragraph(s, style='List Number')
    for run in p.runs:
        run.font.size = Pt(10)

add_para('管理员后台可查看全部退款单，并通过 /api/refunds/{id}/arbitrate 接口做出最终仲裁决定。')

# 4.8 求购广场模块
doc.add_heading('4.8 求购广场模块', level=2)
add_para('求购广场是"反向交易"功能——买家主动发布求购需求，卖家看到后可主动联系。求购信息包含：期望商品描述、预算范围、分类标签、联系方式。求购列表支持关键词搜索和分类筛选。发布者可随时取消求购。')

# 4.9 校园大使模块
doc.add_heading('4.9 校园大使模块', level=2)
add_para('校园大使是平台内容质量的第一道防线，负责审核所有新发布商品。')
add_para('校园大使工作流程：', bold=True)
ambassador_steps = [
    '学生通过 /ambassador 页面提交申请（姓名/学号/学院/宿舍楼栋/社群资源/申请理由）；',
    '管理员在后台审核申请 → 审批通过后该用户角色变更为 CAMPUS_AMBASSADOR，分配工号；',
    '大使登录后台 → 进入"商品审核"页面 → 查看待审核商品列表（status=1）；',
    '审核通过：商品状态 1→2（在售），设置30天过期时间；',
    '审核驳回：商品状态 1→5（审核失败），需填写驳回原因；',
    '大使可查看自己的审核记录和推广数据统计（总审核量/通过率/今日审核量等）。',
]
for s in ambassador_steps:
    p = doc.add_paragraph(s, style='List Number')
    for run in p.runs:
        run.font.size = Pt(10)

# 4.10 举报管理模块
doc.add_heading('4.10 举报管理模块', level=2)
add_para('用户可举报违规商品或不当用户行为。举报信息包含：举报类型（商品/用户）、被举报对象ID、举报原因、详细描述。管理员在后台可查看举报列表，选择"已处理"或"驳回"，并填写处理备注。')

# 4.11 个性化推荐模块
doc.add_heading('4.11 个性化推荐模块', level=2)
add_para('系统基于用户行为数据提供智能推荐：')
add_para('• 个性化推荐（GET /api/recommend）：基于用户浏览(1)、收藏(2)、购买(3)、搜索(4)、点击推荐(5)等行为，使用协同过滤算法推荐商品，默认返回15条。')
add_para('• 关联推荐（GET /api/recommend/related）：基于商品相似度计算"看了这个的人也在看"，默认返回6条。')
add_para('• 行为记录（POST /api/recommend/feedback）：记录用户对推荐结果的反馈（点击/忽略），用于优化推荐模型的 CTR 指标。')
add_para('用户行为日志存储在 user_behavior_log 表，包含用户ID、商品ID、行为类型（1-5）、时间戳。')

# 4.12 管理后台模块
doc.add_heading('4.12 管理后台模块', level=2)
add_para('管理后台通过 /admin 路径访问，根据角色显示不同功能面板：')
add_table(
    ['功能', '路径', '权限'],
    [
        ['数据看板', '/admin/dashboard', '管理员/校园大使'],
        ['用户管理', '/admin/users', '管理员专属'],
        ['商品管理', '/admin/goods', '管理员/校园大使'],
        ['订单管理', '/admin/orders', '管理员/校园大使'],
        ['商品审核', '/admin/review', '校园大使专属'],
        ['推广统计', '/admin/promotion-stats', '校园大使专属'],
        ['举报管理', '/admin/reports', '管理员专属'],
        ['大使申请审核', '/admin/ambassador-applications', '管理员专属'],
    ]
)
add_para('数据看板统计指标：总用户数、总商品数、在售商品数、总订单数、待处理退款数。用户管理支持：搜索（学号/昵称）、封禁/解封、信用分调整、提升/降级校园大使、删除用户。')

# ══════════════════════════════════════════
# 五、数据库设计
# ══════════════════════════════════════════
doc.add_heading('五、数据库设计', level=1)

add_para('系统使用 SQL Server 数据库，包含 campus_secondhand 库，共 11 张核心数据表：', bold=True)
add_table(
    ['表名', '说明', '核心字段'],
    [
        ['users', '用户表', 'id, student_id, phone, nickname, avatar, password(MD5), role, credit_score, status, dormitory, worker_id, token_version, email, gender, real_name, ban_expire_time'],
        ['goods', '商品表', 'id, seller_id, title, description, price, original_price, category, images(JSON), status, video, isbn, location, tags, view_count, collect_count, chat_count, expire_time, reviewer_id, review_reason'],
        ['orders', '订单表', 'id, buyer_id, seller_id, goods_id, amount, status, pickup_code, remark, cancel_reason, pay_time, pickup_time, confirm_time, cancel_time'],
        ['chat_messages', '聊天消息表', 'id, session_id, sender_id, receiver_id, content, type, is_read, created_at'],
        ['chat_sessions', '聊天会话表', 'id, goods_id, initiator_id, participant_id, last_message, last_message_time'],
        ['favorites', '收藏表', 'id, user_id, goods_id, created_at'],
        ['reviews', '评价表', 'id, order_id, reviewer_id, target_user_id, rating(1-5), content'],
        ['refunds', '退款表', 'id, order_id, user_id, reason, images(JSON), status, handler_note'],
        ['buy_requests', '求购表', 'id, user_id, title, description, budget_min, budget_max, category, status'],
        ['reports', '举报表', 'id, reporter_id, target_type, target_id, reason, description, status, handler_note'],
        ['ambassador_applications', '大使申请表', 'id, user_id, name, student_id, phone, department, dormitory, community_resource, reason, status, reviewer_id, review_note'],
        ['user_behavior_log', '用户行为日志表', 'id, user_id, goods_id, behavior_type(1-5), created_at'],
    ]
)

# ══════════════════════════════════════════
# 六、安全性设计
# ══════════════════════════════════════════
doc.add_heading('六、安全性设计', level=1)

add_table(
    ['安全措施', '实现方式'],
    [
        ['身份认证', 'JWT Token（HMAC-SHA384），7天过期，token_version 支持强制下线'],
        ['角色授权', '自定义 @RequireRole 注解 + RoleInterceptor 拦截器，方法级权限控制'],
        ['密码加密', 'MD5 哈希存储（生产环境建议升级为 BCrypt）'],
        ['WebSocket 认证', 'WebSocketAuthInterceptor 在握手阶段校验 JWT Token'],
        ['CORS 配置', 'CorsConfig 限制允许的域名和方法'],
        ['敏感词过滤', 'DFA 字典树算法，sensitive-words.txt 词典（137条），实时过滤聊天/发布内容'],
        ['XSS 防护', '前端 Element Plus 组件自带 XSS 过滤'],
        ['取货码安全', '6位随机数字码，一次性使用，验证后失效'],
        ['文件上传安全', '限制文件类型为图片，MinIO 对象存储隔离'],
        ['SQL 注入防护', 'MyBatis-Plus Lambda 查询构造器从源头避免 SQL 注入'],
    ]
)

# ══════════════════════════════════════════
# 七、部署架构
# ══════════════════════════════════════════
doc.add_heading('七、部署架构', level=1)

add_para('7.1 服务组件', bold=True)
add_table(
    ['组件', '端口', '说明'],
    [
        ['Spring Boot 后端', '8088', 'REST API + WebSocket 服务'],
        ['Vue 3 前端 (Vite Dev)', '5173', '开发环境 HMR 热更新'],
        ['SQL Server', '1433', '主数据库'],
        ['Redis', '6379', '缓存 + 会话管理'],
        ['MinIO', '9000 (API) / 9001 (Console)', '对象存储'],
    ]
)

add_para('7.2 环境配置', bold=True)
add_para('开发环境：后端通过 Maven 直接运行（mvn spring-boot:run），前端通过 Vite 开发服务器运行（npm run dev），MinIO 本地运行 minio.exe。')
add_para('生产环境建议：前端构建为静态文件（npm run build → dist/），部署到 Nginx；后端打包为 JAR 运行（java -jar）；数据库和缓存使用云服务；MinIO 独立部署或使用云对象存储替代。')

# ══════════════════════════════════════════
# 八、实现步骤
# ══════════════════════════════════════════
doc.add_heading('八、实现步骤', level=1)

add_para('以下是该系统的完整实现步骤，从环境搭建到功能上线，共分为 7 个阶段：', bold=True)

# Step 1
doc.add_heading('阶段一：环境搭建与项目初始化', level=2)
add_table(
    ['步骤', '任务内容', '关键产出'],
    [
        ['1.1', '安装开发环境：JDK 17 + Maven 3.9 + Node.js 18 + SQL Server + Redis + MinIO', '开发环境就绪'],
        ['1.2', '使用 Spring Initializr 创建 Spring Boot 3.2 项目，引入 Web/MyBatis-Plus/Redis/Mail/MinIO/JWT 依赖', 'pom.xml'],
        ['1.3', '使用 Vite 创建 Vue 3 项目，安装 Element Plus/Pinia/Vue Router/Axios', 'package.json'],
        ['1.4', '配置 application.yml（数据源/Redis/JWT/MinIO/邮件），编写 application-template.yml 模板', '配置文件'],
        ['1.5', '配置 CORS 跨域（CorsConfig），允许前端开发服务器访问', 'CorsConfig.java'],
        ['1.6', '编写统一响应类 Result<T> 和全局异常处理器 GlobalExceptionHandler', '公共类'],
    ]
)

# Step 2
doc.add_heading('阶段二：安全认证体系', level=2)
add_table(
    ['步骤', '任务内容', '关键产出'],
    [
        ['2.1', '设计 User 实体类，包含 student_id / phone / nickname / password / role / credit_score 等字段', 'User.java'],
        ['2.2', '编写 init.sql：创建 users 表，预设 admin/20240001 两个种子账号', 'init.sql'],
        ['2.3', '实现 JwtTokenProvider：生成/解析/验证 Token，HMAC-SHA384 签名，7天过期', 'JwtTokenProvider.java'],
        ['2.4', '实现 JwtAuthenticationFilter：从请求头 Authorization: Bearer xxx 提取 Token，注入 UserContext', 'JwtAuthenticationFilter.java'],
        ['2.5', '实现 UserContext 线程变量工具类，存储当前请求的用户ID和角色', 'UserContext.java'],
        ['2.6', '实现 @RequireRole 注解 + RoleInterceptor 拦截器，支持方法级角色校验', '@RequireRole + RoleInterceptor'],
        ['2.7', '实现注册/登录/个人资料接口（UserController + UserService）', 'UserController.java'],
        ['2.8', '实现邮件验证码发送和密码重置接口（MailService + VerificationCodeService）', 'MailService.java'],
    ]
)

# Step 3
doc.add_heading('阶段三：商品与交易核心', level=2)
add_table(
    ['步骤', '任务内容', '关键产出'],
    [
        ['3.1', '设计 Goods 实体类，支持多图(JSON)/视频/分类/成色/地点等丰富字段', 'Goods.java'],
        ['3.2', '实现商品 CRUD 接口（GoodsController + GoodsService）：发布/编辑/删除/下架/刷新', 'GoodsController.java'],
        ['3.3', '实现商品搜索接口：关键词模糊匹配 + 分类筛选 + 价格排序 + 分页', 'GoodsQueryDTO'],
        ['3.4', '设计 Order 实体类，含 status/pickup_code 等字段', 'Order.java'],
        ['3.5', '实现订单完整流程：下单→支付→生成取货码→验证→确认收货→取消（OrderController）', 'OrderController.java'],
        ['3.6', '实现取货码机制：generatePickupCode 生成6位随机数，verifyPickup 校验并变更状态', 'OrderService'],
    ]
)

# Step 4
doc.add_heading('阶段四：即时通讯', level=2)
add_table(
    ['步骤', '任务内容', '关键产出'],
    [
        ['4.1', '配置 WebSocket（WebSocketConfig），注册 ChatWebSocketHandler 到 /ws/chat', 'WebSocketConfig.java'],
        ['4.2', '实现 WebSocketAuthInterceptor：握手时校验 JWT Token', 'WebSocketAuthInterceptor.java'],
        ['4.3', '实现 ChatWebSocketHandler：连接管理（ConcurrentHashMap<userId, session>）、消息转发', 'ChatWebSocketHandler.java'],
        ['4.4', '设计 ChatSession + ChatMessage 实体类，实现 ChatSessionService', '会话/消息实体'],
        ['4.5', '实现 ChatController REST API：获取会话列表、获取历史消息、创建会话、上传聊天图片', 'ChatController.java'],
        ['4.6', '实现 DFA 敏感词过滤器（SensitiveWordFilter），加载 sensitive-words.txt', 'SensitiveWordFilter.java'],
    ]
)

# Step 5
doc.add_heading('阶段五：辅助功能模块', level=2)
add_table(
    ['步骤', '任务内容', '关键产出'],
    [
        ['5.1', '实现收藏模块：添加/取消/检查/列表（FavoriteController + FavoriteService）', 'FavoriteController.java'],
        ['5.2', '实现评价模块：提交评价/订单评价/用户评价列表（ReviewController + ReviewService）', 'ReviewController.java'],
        ['5.3', '实现退款模块：申请/同意/拒绝/仲裁（RefundController + RefundService）', 'RefundController.java'],
        ['5.4', '实现求购广场：发布/取消/搜索/我的求购/详情（BuyRequestController）', 'BuyRequestController.java'],
        ['5.5', '实现举报模块：创建举报/列表/处理（ReportController + ReportService）', 'ReportController.java'],
    ]
)

# Step 6
doc.add_heading('阶段六：校园大使与管理后台', level=2)
add_table(
    ['步骤', '任务内容', '关键产出'],
    [
        ['6.1', '设计 AmbassadorApplication 实体，实现校园大使申请公开接口', 'AmbassadorApplication.java'],
        ['6.2', '实现 AmbassadorController：提交申请/查询状态/商品审核(通过-驳回)/审核记录/推广统计', 'AmbassadorController.java'],
        ['6.3', '实现 AdminController：用户管理/商品管理/订单管理/退款管理/大使申请审核/数据看板', 'AdminController.java'],
        ['6.4', '前端 AdminLayout + 权限路由守卫：根据 role 判断 canAccessAdmin / isAdmin / isAmbassador', 'router/index.js'],
        ['6.5', '前端 Dashboard 数据看板：总用户数/总商品数/总订单数/在售商品/待处理退款', 'Dashboard.vue'],
        ['6.6', '前端管理页面：用户列表/商品列表/订单列表/举报管理/大使申请审核', 'admin/ 目录下页面'],
    ]
)

# Step 7
doc.add_heading('阶段七：推荐系统与上线优化', level=2)
add_table(
    ['步骤', '任务内容', '关键产出'],
    [
        ['7.1', '设计 UserBehaviorLog 实体，记录浏览/收藏/购买/搜索/点击行为', 'UserBehaviorLog.java'],
        ['7.2', '实现 RecommendService：协同过滤推荐 + 基于物品相似度的关联推荐', 'RecommendService.java'],
        ['7.3', '实现 RecommendController：个性化推荐/关联推荐/行为反馈接口', 'RecommendController.java'],
        ['7.4', '实现 MinIO 文件服务（FileService + FileController）：上传/图片代理/头像/聊天图片', 'FileController.java'],
        ['7.5', '实现定时任务（SchedulingConfig）：商品30天自动过期、信用分定期更新', 'SchedulingConfig.java'],
        ['7.6', '前端适配：所有页面对接 API、路由守卫完善、主题样式统一', 'admin/src/views/'],
        ['7.7', '数据库种子数据：seed_100_goods.sql（100条商品）、seed_buy_requests.sql（45条求购）', '种子SQL文件'],
        ['7.8', '测试与优化：API 联调测试、性能优化（Redis 缓存热点数据）、安全加固', '—'],
    ]
)

doc.add_paragraph()
add_para('── 文档结束 ──', bold=True, size=10, color=(0x99, 0x99, 0x99), align=WD_ALIGN_PARAGRAPH.CENTER)

# 保存文件
output_path = 'e:/校园二手交易平台_功能概述与实现步骤.docx'
doc.save(output_path)
print(f'文档已生成：{output_path}')
